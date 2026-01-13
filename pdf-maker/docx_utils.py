from docx import Document

def generate_docx(text):
    doc = Document()
    for line in text.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)

        if line.strip().upper().startswith(("UNIT", "SECTION")):
            run.bold = True

    file_path = "output.docx"
    doc.save(file_path)

    with open(file_path, "rb") as f:
        return f.read()
